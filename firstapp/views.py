from django.shortcuts import render
from django.http import HttpResponse
from .forms import UserForm
from docxtpl import DocxTemplate
from docx import Document
from django.core.mail import send_mail
from django.core.mail import EmailMultiAlternatives
from docx.shared import Inches

def index(request):
    if request.method == "POST":
        form = UserForm(request.POST)
        if form.is_valid():
            name = request.POST.get("name")
            surname = request.POST.get("surname")
            email = request.POST.get("email")
            check1 = form.cleaned_data['check1']
            check2 = form.cleaned_data['check2']
            photos = form.cleaned_data.get('photos')
            doc1 = DocxTemplate(r"C:\Users\Валерия\hello\DOC.docx")
            doc3 = DocxTemplate(r"C:\Users\Валерия\hello\TZ.docx")
            context = {'name': name, 'surname': surname,'email':email}
            doc1.render(context)
            doc3.render(context)
            doc1.save("C:/Users/Валерия/hello/DOC1CHANGE.docx")
            doc3.save("C:/Users/Валерия/hello/DOC3CHANGE.docx")
            t1 = Document(r"C:\Users\Валерия\hello\DOC1CHANGE.docx")
            t3 = Document(r"C:\Users\Валерия\hello\DOC3CHANGE.docx")
            t2 = Document(r"C:\Users\Валерия\hello\PRAVA.docx")
            if check1 == True:
                for p in t2.paragraphs:
                    t1.add_paragraph(p.text, p.style)
            if check2 == True:
                for p in t3.paragraphs:
                    t1.add_paragraph(p.text, p.style)
            for i in photos:
                if i == "front":
                    t1.add_picture("C:/Users/Валерия/hello/front.jpg", width=Inches(4.0))
                if i == "1floor":
                    t1.add_picture("C:/Users/Валерия/hello/1floor.png", width=Inches(4.0))
                if i == "2floor":
                    t1.add_picture("C:/Users/Валерия/hello/2floor.png", width=Inches(4.0))
            t1.save("C:/Users/Валерия/hello/FINAL.docx")
            msg = EmailMultiAlternatives("Документы", "Ваш пакет документов", "vorontsova000@gmail.com", [email])
            msg.attach_file('C:/Users/Валерия/hello/FINAL.docx')
            msg.send()
            #send_mail('Subject here', 'Here is the message.', 'from@example.com',
                      #[email], fail_silently=False)
            return HttpResponse("<h2>Ваш пакет документов успешно отправлен на {0}</h2>".format(email))
    else:
        userform = UserForm()
        return render(request, "index.html", {"form": userform})
    